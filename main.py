import docx
import telebot
import requests

url = 'https://адрес-сайта/путь-до-файла.docx'
headers = {'user-agent': 'Mozilla/5.0'}

# Скачиваем документ с сайта
# Прикидываемся браузером, иначе получим ошибку 403
r = requests.get(url, headers = headers)
with open("./download/имя-файла.docx", 'wb') as f:
    f.write(r.content)

# Инициализация переменных
flag = False                                                                # flag: устанаилвается в True, если найдены изменения
bot = telebot.TeleBot('ID-бота')                                            # bot:  инициализация бота с ID, полученным от @BotFather
doc1 = docx.Document("./исходный-файл.docx")                                # doc1: исходный файл для сравнения
doc2 = docx.Document("./download/файл-для-сравнения.docx")                  # doc2: скачанный файл с сайта
doc1paragraphs = []
doc2paragraphs = []
 
# Начинаем анализ текста в разрезе параграфов
for paragraph in doc1.paragraphs:
    doc1paragraphs.append(paragraph.text)
for paragraph in doc2.paragraphs:
    doc2paragraphs.append(paragraph.text)

# Для верной обработки кириллицы задаем кодировку UTF-16
sourceFile = open('./log.txt', 'w', encoding = 'utf-16')

# Проверям параграфы на соответствие, если нет - запись в лог и оповещение
for i in doc1paragraphs:
    if i in doc2paragraphs:
        print(f"[MATCH] {i}", file = sourceFile)
    else:
        print(f"[NO MATCH] {i}", file = sourceFile)
        flag = True

sourceFile.close()

# Если нашли несоответствие - пишем в указанный чат. Первый аргумент - ID чата
if flag:
    bot.send_message("ID-чата", url + " - есть изменения в документе")
else:
    bot.send_message("ID-чата", url + " - нет изменений в документе")